import os

import openpyxl
import uuid
from django.db.models import QuerySet
from docx import Document

from src.apps.core.service import upload_file
from src.apps.neuro.answers import make_answer
from src.quick_check.settings import MEDIA_ROOT


def parse_questions(file_path):
    result = []
    workbook = openpyxl.open(file_path)
    sheet = workbook.worksheets[0]
    question_text_col = 'A'
    question_answer_col = 'B'
    num = 1
    for row in range(2, sheet.max_row + 1):
        text = sheet[question_text_col + str(row)].value
        answer = sheet[question_answer_col + str(row)].value
        if text and answer:
            result.append((num, text, str(answer).lower()))
            num += 1
    workbook.close()
    return result


def parse_text(file_path: str):
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
    else:
        # parse as txt
        with open(file_path, 'r') as file:
            text = file.read()
            return text


def extract_images_and_text_as_list(file_path):
    doc = Document(file_path)
    pupils = []
    images = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            pupils.append(str(paragraph.text).lower())

    for shape in doc.inline_shapes:
        content_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        content_type = doc.part.related_parts[content_id].content_type
        if not content_type.startswith('image'):
            continue
        # img_format = str(doc.part.related_parts[content_id].partname).strip('.')[-1]
        img_format = 'jpeg'
        img_name = '{}.{}'.format(uuid.uuid4(), img_format)
        image_data = doc.part.related_parts[content_id]._blob
        image_file_path = os.path.join(MEDIA_ROOT, img_name)
        with open(image_file_path, 'wb') as file:
            file.write(image_data)
        images.append(upload_file(image_file_path))
    return pupils, images


def get_question_results(file):
    """Отправление файла к машинке и получение ответов"""
    return []


def get_homework_result(questions: QuerySet, file):
    """Получение ответов от машинки (пока заглушка)"""
    print('questions count {}'.format(questions.count()))
    answers = make_answer(questions.count(), file)
    print('answers')
    print(answers)
    has_answer_questions = []
    result = []
    for num in range(1, max(answers.answers.keys())):
        try:
            ans = (''.join(answers.answers.get(num, []))).lower()
            print('question {}, ans {}'.format(num, ans))
            question = questions.get(num=num)
            correct_answer = question.answer.lower()
            if ans and (ans in correct_answer or correct_answer in ans):
                ans = question.answer
            check = question.answer.lower() == ans
            result.append((question, ans, check))
            has_answer_questions.append(question)
        except Exception as e:
            print('get_homework_result error, number {}'.format(num))
            print(e)

    # Вопросы, на которые не было дано ответа
    for question in questions:
        if question not in has_answer_questions:
            result.append((question, '', False))
    return result, answers.data


class SequenceAlignment(object):
    def __init__(self, x, y):
        self.x = x
        self.y = y
        self.solution = []

    delta = lambda self, x, y, i, j: 1 if x[i] != y[j] else 0

    def find_solution(self, temp, m, n):
        if m == 0 and n == 0:
            return

        insert = temp[m][n - 1] + 1 if n != 0 else float("inf")
        align = (
            temp[m - 1][n - 1] + self.delta(self.x, self.y, m - 1, n - 1)
            if m != 0 and n != 0
            else float("inf")
        )
        delete = temp[m - 1][n] + 1 if m != 0 else float("inf")

        best_choice = min(insert, align, delete)

        if best_choice == insert:
            self.solution.append(True)
            return self.find_solution(temp, m, n - 1)

        elif best_choice == align:
            mistake_flag = self.x[m-1] != self.y[n-1]
            self.solution.append(mistake_flag)
            return self.find_solution(temp, m - 1, n - 1)

        elif best_choice == delete:
            self.solution.append(True)
            return self.find_solution(temp, m - 1, n)

    def alignment(self):
        n = len(self.y)
        m = len(self.x)
        temp = [[0 for _ in range(n + 1)] for _ in range(m + 1)]

        for i in range(1, m + 1):
            temp[i][0] = i

        for j in range(1, n + 1):
            temp[0][j] = j

        for i in range(1, m + 1):
            for j in range(1, n + 1):
                temp[i][j] = min(
                    temp[i - 1][j - 1] + self.delta(self.x, self.y, i - 1, j - 1),
                    temp[i - 1][j] + 1,
                    temp[i][j - 1] + 1,
                )

        self.find_solution(temp, m, n)

        return temp[m][n], self.solution[::-1]
